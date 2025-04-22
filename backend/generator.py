import os
from dotenv import load_dotenv
import google.generativeai as genai
from backend.prompts import resume_prompt, cover_letter_prompt
from docx import Document
import re
import json

# Load environment variables
load_dotenv()

# Configure the Google API key
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
genai.configure(api_key=GOOGLE_API_KEY)

# Set the default model for Gemini
DEF_MODEL = "models/gemini-2.0-flash-exp"


def get_model(model_name=DEF_MODEL):
    return genai.GenerativeModel(model_name)


def generate_resume(name, experience, skills, job_desc, education, save_doc=True):
    prompt = resume_prompt.format(name=name, experience=experience, skills=skills, job_desc=job_desc, education=education)
    model = get_model()
    response = model.generate_content(prompt)
    resume_text = response.text

    if save_doc:
        try:
            json.loads(resume_text)  # test if it's valid JSON
            filename = f"{name.replace(' ', '_')}_Resume.docx"
            return save_json_resume_to_docx(resume_text, filename)
        except json.JSONDecodeError:
            filename = f"{name.replace(' ', '_')}_Resume.docx"
            return save_markdown_like_text_to_docx(resume_text, filename)

    return resume_text


def generate_cover_letter(name, company, job_desc, save_doc=True):
    # Format the cover letter prompt
    prompt = cover_letter_prompt.format(
        name=name,
        company=company,
        job_desc=job_desc
    )

    # Get the Gemini model
    model = get_model()

    # Generate the response
    response = model.generate_content(prompt)
    cover_letter_text = response.text

    if save_doc:
        filename = f"{name.replace(' ', '_')}_CoverLetter.docx"
        save_markdown_like_text_to_docx(cover_letter_text, filename)

    return cover_letter_text

def save_markdown_like_text_to_docx(text, filename="output.docx"):
    # Remove Markdown code block wrappers (e.g., ```markdown and ```)
    text = re.sub(r'^```markdown\n', '', text)
    text = re.sub(r'\n```$', '', text)

    # Create a new Word document
    doc = Document()

    # Process each line
    for line in text.split('\n'):
        # Check for bold markdown style (e.g., **bold**)
        if line.startswith("**") and line.endswith("**"):
            # Heading (bold text)
            doc.add_paragraph(line[2:-2], style="Heading 1")  # You can change Heading 1 to Heading 2 for smaller headings
        # Check for bullet points
        elif line.startswith("* ") or line.startswith("- "):
            # Bullet points
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            # Regular text (non-formatted)
            doc.add_paragraph(line)

    # Save the document
    output_path = os.path.join("generated_docs", filename)
    os.makedirs("generated_docs", exist_ok=True)
    doc.save(output_path)
    print(f"Saved Word document to {output_path}")
    return output_path

    ### The save_markdown_like_text_to_docx function was generated with ChatGPT. Prompt: "Can you write a Python function that takes markdown-style text
    ### (with **bold**, bullet points, and ```markdown blocks), and converts it into a nicely formatted .docx Word
    ### document using python-docx? It should handle headings, bullets, and save to a folder."

def save_json_resume_to_docx(json_text, filename="resume.docx"):
    data = json.loads(json_text)
    doc = Document()

    # Add Name as Heading
    doc.add_heading(data.get("personal_information", {}).get("name", "Name"), level=1)

    # Add Summary
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(data.get("summary", ""))

    # Experience
    doc.add_heading("Experience", level=2)
    for job in data.get("experience", []):
        doc.add_paragraph(f"{job['title']} at {job['company']} ({job['dates']})", style="List Bullet")
        for resp in job.get("responsibilities", []):
            doc.add_paragraph(resp, style="List Bullet 2")

    # Education
    doc.add_heading("Education", level=2)
    for edu in data.get("education", []):
        doc.add_paragraph(f"{edu['degree']} - {edu['university']} ({edu['dates']})", style="List Bullet")

    # Skills
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(data.get("skills", [])))

    # Save file
    output_path = os.path.join("generated_docs", filename)
    os.makedirs("generated_docs", exist_ok=True)
    doc.save(output_path)
    return output_path

# Usage examples
if __name__ == '__main__':
    # Example for generating a resume
    name = "Jane Doe"
    experience = "5 years of experience as data scientist at Microsoft"
    skills = "Python, R, Machine Learning"
    education = "Bachelor in Data Science from Penn State, Masters in Business Intelligence and Analytics from SJU"
    job_desc = "Machine Learning engineer"
    print(generate_resume(name, experience, skills, education, job_desc))

    # Example for generating a cover letter
    name = "John Doe"
    company = "Microsoft"
    job_desc = "Data Base Engineer"
    print(generate_cover_letter(name, company, job_desc))

